# Copyright (c) 2025 Cobus Nel
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.
"""
Render dkit doc cannonical format to Micosoft Word docx format using docx
"""
from . import document as doc
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from .docx_helper import add_hyperlink, create_codeblock_style, DocxConfig
import functools
HEADING_COUNTER = 0


class DocxRenderer:
    """Render document elements to docx"""

    def __init__(self, document: doc.Document, allow_soft_breaks: bool = False,
                 config: DocxConfig = None, template: str = None):
        if config is None:
            config = DocxConfig()
        self.allow_soft_breaks = allow_soft_breaks  # allow breaks in a paragraph
        self.doc = document
        self.config = config
        if template is None:
            self.xdoc = docx.Document()
        else:
            self.xdoc = docx.Document(template)
        # Create custom style
        create_codeblock_style(self.xdoc)

        # State variables
        self.current_paragraph = None
        self.current_style = self.config.sty_normal
        self.list_level = None  # used by lists to determine level
        self.list_type = None     # used by lists

    def set_properties(self):
        """set document properties"""
        self.xdoc.core_properties.author = self.doc.author
        self.xdoc.core_properties.title = self.doc.title
        self.xdoc.core_properties.subject = self.doc.sub_title

    def make_elements(self, elements):
        for element in elements:
            self.make(element)

    @functools.singledispatchmethod
    def make(self, element):
        raise TypeError(f"Unsupported data type: {type(element)}")

    @make.register(doc.Code)
    def make_code(self, element: doc.Str):
        run = self.current_paragraph.add_run(element.content)
        return run

    @make.register(doc.Str)
    def make_str(self, element: doc.Str):
        return self.current_paragraph.add_run(element.text)

    @make.register(doc.Bold)
    def make_bold(self, element: doc.Bold):
        for elem in element.text:
            run = self.make(elem)
            run.font.bold = True

    @make.register(doc.Emph)
    def make_emph(self, element: doc.Emph):
        for elem in element.text:
            run = self.make(elem)
            run.font.italic = True

    @make.register(doc.List)
    def make_list(self, element: doc.List):
        buf = self.current_style
        self.list_level = element.depth + 1
        if element.ordered is True:
            self.list_type = self.config.sty_number_list
        else:
            self.list_type = self.config.sty_bullet
        self.current_style = self._get_bullet_style()
        for elem in element.content:
            self.make(elem)
        self.current_style = buf

    @make.register(doc.ListItem)
    def make_list_item(self, element: doc.ListItem):
        self._add_paragraph()
        self.make_elements(element.content)

    @make.register(doc.Block)
    def make_block(self, element: doc.Block):
        self.make_elements(element.content)

    def _add_paragraph(self, text="", style=None):
        if style is not None:
            self.current_style = style
        self.current_paragraph = self.xdoc.add_paragraph(
            text=text,
            style=self.current_style
        )
        return self.current_paragraph

    def _get_bullet_style(self):
        if self.list_level == 1:
            return f"List {self.list_type}"
        elif self.list_level > 1:
            return f"List {self.list_type} {self.list_level}"
        else:
            raise ValueError(f"bullet level {self.list_level} is invalid")

    @make.register(doc.Heading)
    def make_heading(self, element: doc.Heading):
        self.current_paragraph = self.xdoc.add_heading(
            "",
            element.level
        )
        self.make_elements(element.content)

    @make.register(doc.Paragraph)
    def make_paragraph(self, element: doc.Paragraph):
        self.current_paragraph = self._add_paragraph()
        self.make_elements(element.content)

    @make.register(doc.LineBreak)
    def make_line_break(self, element: doc.LineBreak):
        self.current_paragraph = self.xdoc.add_paragraph()
        """
        run = self.xdoc.add_paragraph().add_run()
        run.add_break()
        """

    @make.register(doc.Link)
    def make_hyperlink(self, element: doc.Link):
        add_hyperlink(
            self.current_paragraph,
            self._reduce_to_text(element.content),
            element.target
        )

    def _reduce_to_text(self, elements):
        """reduce to text, for use by hyperlinks etc"""
        try:
            return " ".join(i.text for i in elements)
        except TypeError:
            raise TypeError(f"Only Str elements expected here: {elements}")

    @make.register(doc.SoftBreak)
    def make_soft_break(self, element: doc.SoftBreak):
        if self.allow_soft_breaks:
            run = self.current_paragraph.add_run()
            run.add_break()

    @make.register(doc.PageBreak)
    def make_page_break(self, element: doc.PageBreak):
        self.xdoc.add_page_break()

    @make.register(doc.CodeBlock)
    def make_code_block(self, element: doc.CodeBlock):
        buffer = self.current_style  # Save for later"
        self._add_paragraph(
            element.content,
            style=self.config.sty_code_block
        )
        self._add_paragraph(style=buffer)

    @make.register(doc.BlockQuote)
    def make_block_quote(self, element: doc.BlockQuote):
        buffer = self.current_style  # Save for later"
        self.current_style = self.config.sty_quote
        self.make_elements(element.content)
        self.current_style = buffer

    @make.register(doc.Image)
    def make_image(self, element: doc.Image):
        para = self._add_paragraph()
        run = para.add_run()
        run.add_picture(
            element.source,
            width=Cm(element.width) if element.width else None,
            height=Cm(element.height) if element.height else None
        )
        para.alignment = self._translate_alignment(element.align)

    def _translate_alignment(self, align: str):
        _align = align.lower()
        match _align:
            case "left":
                return WD_ALIGN_PARAGRAPH.LEFT
            case "right":
                return WD_ALIGN_PARAGRAPH.RIGHT
            case "center":
                return WD_ALIGN_PARAGRAPH.CENTER
        return WD_ALIGN_PARAGRAPH.CENTER

    def _translate_table_alignment(self, align: str):
        _align = align.lower()
        match _align:
            case "left":
                return WD_TABLE_ALIGNMENT.LEFT
            case "right":
                return WD_TABLE_ALIGNMENT.RIGHT
            case "center":
                return WD_TABLE_ALIGNMENT.CENTER
        return WD_TABLE_ALIGNMENT.CENTER

    @make.register(doc.Table)
    def make_table(self, element: doc.Table):
        rows = len(element.data) + 1
        cols = len(element.columns)
        table = self.xdoc.add_table(rows, cols, style=self.config.sty_table)
        table.autofit = False
        table.alignment = self._translate_table_alignment(element.align)

        # set table column widths
        for i, col in enumerate(element.columns):
            table.columns[i].width = Cm(col.width)

        # set headers
        headers = table.rows[0].cells
        for i, col in enumerate(element.columns):
            header = headers[i]
            header.text = col.title
            header.width = Cm(col.width)
            for p in header.paragraphs:
                p.alignment = self._translate_alignment(col.heading_align)

        # set data
        for i, data_row in enumerate(element.data):
            row = table.rows[i + 1]
            for j, col in enumerate(element.columns):
                cell = row.cells[j]
                cell.text = str(data_row[col.name])
                cell.width = col.width
                for p in cell.paragraphs:
                    p.alignment = self._translate_alignment(col.align)

    def render(self, file_name: str):
        self.set_properties()
        self.make_elements(self.doc.elements)
        self.xdoc.save(file_name)
